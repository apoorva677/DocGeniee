import os
from xhtml2pdf import pisa
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from backend.models.document_model import Document

class ExportEngine:
    """
    ExportEngine handles exporting documents to DOCX and styled PDF formats.
    """

    def export_docx(self, document: Document, path: str):
        doc = DocxDocument()
        doc.add_heading(document.title, 0)
        
        for section in document.sections:
            doc.add_heading(section.heading, level=min(section.level, 9))
            for block in section.blocks:
                p = doc.add_paragraph()
                if block.type == "list":
                    p.style = 'List Bullet'
                run = p.add_run(block.content)
                
                if hasattr(block, 'style') and block.style:
                    if 'font-size' in block.style:
                        size_str = block.style['font-size'].replace('pt', '').strip()
                        if size_str.isdigit(): 
                            run.font.size = Pt(int(size_str))
                    if 'font-family' in block.style:
                        run.font.name = block.style['font-family']
                    if 'font-weight' in block.style and block.style['font-weight'] in ['bold', '600', '700']:
                        run.bold = True
        doc.save(path)

    def export_pdf(self, document: Document, path: str):
        htmlContent = f"<h1>{document.title}</h1>\n"
        for sec in document.sections:
            hLevel = min((sec.level or 1) + 1, 6)
            htmlContent += f"<h{hLevel}>{sec.heading}</h{hLevel}>\n"
            for block in sec.blocks:
                styleStr = ""
                if hasattr(block, 'style') and block.style:
                    styleStr = ' style="' + '; '.join([f"{k}: {v}" for k, v in block.style.items()]) + '"'
                
                if block.type == "heading":
                    # For inside-block headings if any
                    htmlContent += f"<h{hLevel}{styleStr}>{block.content}</h{hLevel}>\n"
                elif block.type == "paragraph":
                    htmlContent += f"<p{styleStr}>{block.content}</p>\n"
                elif block.type == "list":
                    htmlContent += f"<ul><li{styleStr}>{block.content}</li></ul>\n"
                else:
                    htmlContent += f"<p{styleStr}>{block.content}</p>\n"
                    
        # Wrap into full HTML boilerplate to ensure xhtml2pdf has context
        full_html = f"<html><head><style>body {{ font-family: Helvetica, Arial, sans-serif; }}</style></head><body>{htmlContent}</body></html>"
        with open(path, "w+b") as result_file:
            pisa_status = pisa.CreatePDF(full_html, dest=result_file)